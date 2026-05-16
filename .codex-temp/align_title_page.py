from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


DOC_PATH = Path(r"A:\Proiecte\SwiftBank\documentatie\Licenta Cirimpei Luca.docx")
BACKUP_PATH = DOC_PATH.with_name("Licenta Cirimpei Luca - backup inainte aliniere pagina titlu.docx")
OUT_PATH = DOC_PATH.with_name("Licenta Cirimpei Luca - pagina titlu aliniata.docx")

TITLE_LINES = [
    "SISTEM INFORMATIC DE TIP E-BANKING",
    "CU FUNCȚIONALITĂȚI MODERNE DE SECURITATE,",
    "ANALIZĂ FINANCIARĂ ȘI PROCESARE",
    "A TRANZACȚIILOR",
]


def set_run_font(run, size=14, bold=True):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold


def clear_and_set_title(paragraph):
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.08
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    for index, line in enumerate(TITLE_LINES):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, size=14, bold=True)


def configure_spacer(paragraph, after_points):
    paragraph.clear()
    paragraph.alignment = None
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after_points)
    paragraph.paragraph_format.line_spacing = 1.0


def normalize_second_page_header(document):
    for paragraph in document.paragraphs:
        text = " ".join(paragraph.text.split())
        if text == "ACADEMIA DE STUDII ECONOMICE DIN BUCUREȘTI":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(8)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                set_run_font(run, size=13, bold=True)
        elif text == "Facultatea de Cibernetică, Statistică și Informatică Economică":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                set_run_font(run, size=12, bold=False)
        elif text == "Specializarea: Informatică economică":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                set_run_font(run, size=12, bold=False)


def insert_after(anchor, element):
    parent = anchor.getparent()
    parent.insert(parent.index(anchor) + 1, element)


def main():
    if not DOC_PATH.exists():
        raise FileNotFoundError(DOC_PATH)

    if not BACKUP_PATH.exists():
        shutil.copy2(DOC_PATH, BACKUP_PATH)

    document = Document(DOC_PATH)
    body = document.element.body

    if len(document.tables) < 2:
        raise RuntimeError("Expected two front-matter tables")

    title_paragraph = next(
        (
            paragraph
            for paragraph in document.paragraphs
            if "SISTEM INFORMATIC DE TIP E-BANKING" in paragraph.text
        ),
        None,
    )
    if title_paragraph is None:
        raise RuntimeError("Title paragraph not found")

    city_paragraphs = [
        paragraph for paragraph in document.paragraphs if "București" in paragraph.text
    ]
    if len(city_paragraphs) < 2:
        raise RuntimeError("Second page city/year paragraph not found")
    city_paragraph = city_paragraphs[-1]

    second_table = document.tables[1]
    title_elem = title_paragraph._p
    city_elem = city_paragraph._p
    table_elem = second_table._tbl

    # Rebuild the visual order of page 2 to mirror the first page:
    # header -> spacer -> title -> spacer -> coordinator/graduate block -> spacer -> city/year.
    clear_and_set_title(title_paragraph)
    normalize_second_page_header(document)

    spacer_after_title = OxmlElement("w:p")
    spacer_before_city = OxmlElement("w:p")
    spacer_after_title_paragraph = Paragraph(spacer_after_title, document._body)
    spacer_before_city_paragraph = Paragraph(spacer_before_city, document._body)
    configure_spacer(spacer_after_title_paragraph, 104)
    configure_spacer(spacer_before_city_paragraph, 52)

    for element in (table_elem, city_elem):
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)

    insert_after(title_elem, spacer_after_title)
    insert_after(spacer_after_title, table_elem)
    insert_after(table_elem, spacer_before_city)
    insert_after(spacer_before_city, city_elem)

    city_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    city_paragraph.paragraph_format.space_before = Pt(0)
    city_paragraph.paragraph_format.space_after = Pt(0)
    city_paragraph.paragraph_format.line_spacing = 1.0
    for run in city_paragraph.runs:
        set_run_font(run, size=12, bold=False)

    try:
        document.save(DOC_PATH)
        print(DOC_PATH)
    except PermissionError:
        document.save(OUT_PATH)
        print(OUT_PATH)

    print(BACKUP_PATH)


if __name__ == "__main__":
    main()
