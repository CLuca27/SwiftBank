from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


DOC_PATH = Path(r"A:\Proiecte\SwiftBank\documentatie\Licenta Cirimpei Luca.docx")
BACKUP_PATH = DOC_PATH.with_name("Licenta Cirimpei Luca - backup inainte titlu.docx")
OUT_PATH = DOC_PATH.with_name("Licenta Cirimpei Luca - titlu aranjat.docx")

OLD_TITLE = (
    "SISTEM INFORMATIC DE TIP E-BANKING CU FUNCȚIONALITĂȚI MODERNE DE SECURITATE, "
    "ANALIZĂ FINANCIARĂ SI PROCESARE A TRANZACȚIILOR"
)

TITLE_LINES = [
    "SISTEM INFORMATIC DE TIP E-BANKING",
    "CU FUNCȚIONALITĂȚI MODERNE DE SECURITATE,",
    "ANALIZĂ FINANCIARĂ ȘI PROCESARE",
    "A TRANZACȚIILOR",
]


def set_font(run, size=14, bold=True):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold


def main():
    if not DOC_PATH.exists():
        raise FileNotFoundError(DOC_PATH)

    if not BACKUP_PATH.exists():
        shutil.copy2(DOC_PATH, BACKUP_PATH)

    document = Document(DOC_PATH)
    title_paragraph = None

    for paragraph in document.paragraphs:
        normalized = " ".join(paragraph.text.split())
        if normalized == " ".join(OLD_TITLE.split()):
            title_paragraph = paragraph
            break

    if title_paragraph is None:
        raise RuntimeError("Title paragraph not found")

    title_paragraph.clear()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.line_spacing = 1.08
    title_paragraph.paragraph_format.space_before = Pt(0)
    title_paragraph.paragraph_format.space_after = Pt(0)

    for index, line in enumerate(TITLE_LINES):
        if index:
            title_paragraph.add_run().add_break()
        run = title_paragraph.add_run(line)
        set_font(run, size=14, bold=True)

    try:
        document.save(DOC_PATH)
        saved_path = DOC_PATH
    except PermissionError:
        document.save(OUT_PATH)
        saved_path = OUT_PATH

    print(saved_path)
    print(BACKUP_PATH)


if __name__ == "__main__":
    main()
