from pathlib import Path
import shutil

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOC_PATH = Path(r"A:\Proiecte\SwiftBank\documentatie\Licenta Cirimpei Luca.docx")
BACKUP_PATH = DOC_PATH.with_name("Licenta Cirimpei Luca - backup layout stricat.docx")
OUT_PATH = DOC_PATH.with_name("Licenta Cirimpei Luca - layout reparat.docx")

UNIVERSITY = "ACADEMIA DE STUDII ECONOMICE DIN BUCUREȘTI"
FACULTY = "Facultatea de Cibernetică, Statistică și Informatică Economică"
SPECIALIZATION = "Specializarea: Informatică economică"
WORK_LABEL = "LUCRARE DE LICENȚĂ"
TITLE_LINES = [
    "SISTEM INFORMATIC DE TIP E-BANKING",
    "CU FUNCȚIONALITĂȚI MODERNE DE SECURITATE,",
    "ANALIZĂ FINANCIARĂ ȘI PROCESARE",
    "A TRANZACȚIILOR",
]
COORDINATOR = "Prof. Univ. Dr. Lorena Pocatilu"
GRADUATE = "Cirimpei Luca"


def set_run_font(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold


def set_no_cell_border(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def clear_cell(cell):
    for paragraph in cell.paragraphs:
        paragraph.clear()


def add_paragraph(document, text="", size=12, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, after=0):
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.0
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, size=size, bold=bold)
    return paragraph


def add_spacer(document, after):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.0
    return paragraph


def add_people_block(document):
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(8)
    table.columns[1].width = Cm(8)

    left = table.cell(0, 0)
    right = table.cell(0, 1)
    for cell in (left, right):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_no_cell_border(cell)
        clear_cell(cell)

    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run("Cadrul didactic coordonator:\n")
    set_run_font(run, 12)
    run = p.add_run(COORDINATOR)
    set_run_font(run, 12, bold=True)

    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run("Absolvent:\n")
    set_run_font(run, 12)
    run = p.add_run(GRADUATE)
    set_run_font(run, 12, bold=True)

    return table


def configure_section(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_header(document):
    add_paragraph(document, UNIVERSITY, size=13, bold=True, after=8)
    add_paragraph(document, FACULTY, size=12, after=4)
    p = add_paragraph(document, "", size=12, after=0)
    run = p.add_run("Specializarea: ")
    set_run_font(run, 12)
    run = p.add_run("Informatică economică")
    set_run_font(run, 12, bold=False)


def add_city_year(document):
    p = add_paragraph(document, "", size=12, after=0)
    run = p.add_run("București")
    set_run_font(run, 12)
    run = p.add_run()
    run.add_break()
    run = p.add_run("2026")
    set_run_font(run, 12)


def build_cover(document):
    add_header(document)
    add_spacer(document, 118)
    add_paragraph(document, WORK_LABEL, size=20, bold=True, after=8)
    add_spacer(document, 150)
    add_people_block(document)
    add_spacer(document, 52)
    add_city_year(document)


def build_title_page(document):
    add_header(document)
    add_spacer(document, 108)

    p = add_paragraph(document, "", after=0)
    p.paragraph_format.line_spacing = 1.05
    for index, line in enumerate(TITLE_LINES):
        if index:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, 14, bold=True)

    add_spacer(document, 132)
    add_people_block(document)
    add_spacer(document, 52)
    add_city_year(document)


def main():
    if DOC_PATH.exists() and not BACKUP_PATH.exists():
        shutil.copy2(DOC_PATH, BACKUP_PATH)

    document = Document()
    configure_section(document.sections[0])

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    build_cover(document)
    configure_section(document.add_section(WD_SECTION_START.NEW_PAGE))
    build_title_page(document)

    try:
        document.save(DOC_PATH)
        saved_path = DOC_PATH
    except PermissionError:
        document.save(OUT_PATH)
        saved_path = OUT_PATH

    print(saved_path)
    print(BACKUP_PATH)


if __name__ == "__main__":
    main()
