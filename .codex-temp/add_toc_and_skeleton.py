from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOC_DIR = Path(r"A:\Proiecte\SwiftBank\documentatie")
SOURCE = DOC_DIR / "Licenta Cirimpei Luca - layout reparat.docx"
FALLBACK_SOURCE = DOC_DIR / "Licenta Cirimpei Luca.docx"
OUT_PATH = DOC_DIR / "Licenta Cirimpei Luca - schelet cuprins.docx"


SECTIONS = [
    (
        "INTRODUCERE",
        [],
    ),
    (
        "CAPITOLUL 1. DESCRIEREA PROBLEMEI ECONOMICE",
        [
            "1.1. Digitalizarea serviciilor bancare",
            "1.2. Rolul aplicațiilor mobile banking în relația client-bancă",
            "1.3. Securitatea operațiunilor financiare digitale",
            "1.4. Plățile electronice și autentificarea tranzacțiilor",
            "1.5. Analiza comparativă a aplicațiilor bancare existente",
        ],
    ),
    (
        "CAPITOLUL 2. ANALIZA ȘI PROIECTAREA SISTEMULUI INFORMATIC SWIFTBANK",
        [
            "2.1. Specificarea cerințelor sistemului informatic",
            "2.2. Actorii sistemului și cazurile de utilizare",
            "2.3. Diagrama cazurilor de utilizare",
            "2.4. Descrierea cazurilor de utilizare principale",
            "2.5. Modelul conceptual al sistemului",
            "2.6. Proiectarea bazei de date",
            "2.7. Diagrame de secvență pentru funcționalitățile aplicației",
            "2.8. Diagrama de componente și arhitectura sistemului",
        ],
    ),
    (
        "CAPITOLUL 3. IMPLEMENTAREA APLICAȚIEI INFORMATICE",
        [
            "3.1. Tehnologii și produse software utilizate",
            "3.2. Implementarea backend-ului",
            "3.3. Implementarea aplicației mobile Android",
            "3.4. Implementarea mecanismelor de autentificare și securitate",
            "3.5. Implementarea funcționalităților bancare",
            "3.6. Implementarea simulatorului de plăți cu cardul",
            "3.7. Notificări, sincronizare în timp real și statistici",
            "3.8. Prezentarea aplicației și scenarii de utilizare",
            "3.9. Limitări și direcții viitoare de dezvoltare",
        ],
    ),
    (
        "CONCLUZII",
        [],
    ),
    (
        "BIBLIOGRAFIE",
        [],
    ),
    (
        "ANEXE",
        [
            "A.1. Fragmente reprezentative de cod sursă",
            "A.2. Diagrame UML suplimentare",
            "A.3. Capturi suplimentare ale aplicației",
        ],
    ),
]


def set_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def configure_document(document):
    for section in document.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    heading1 = document.styles["Heading 1"]
    heading1.font.name = "Times New Roman"
    heading1._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    heading1.font.size = Pt(14)
    heading1.font.bold = True
    heading1.paragraph_format.space_before = Pt(0)
    heading1.paragraph_format.space_after = Pt(12)

    heading2 = document.styles["Heading 2"]
    heading2.font.name = "Times New Roman"
    heading2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    heading2.font.size = Pt(12)
    heading2.font.bold = True
    heading2.paragraph_format.space_before = Pt(10)
    heading2.paragraph_format.space_after = Pt(6)


def set_update_fields_on_open(document):
    settings = document.settings.element
    for child in settings.findall(qn("w:updateFields")):
        settings.remove(child)
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


def add_toc_field(paragraph):
    run = paragraph.add_run()

    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'
    run._r.append(instr_text)

    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char)

    fallback = paragraph.add_run("Cuprinsul se actualizează automat în Word: clic dreapta -> Update Field.")
    set_font(fallback, 12, italic=True)

    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)


def add_toc_page(document):
    document.add_page_break()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(24)
    run = title.add_run("CUPRINS")
    set_font(run, 14, bold=True)

    toc = document.add_paragraph()
    toc.paragraph_format.space_after = Pt(0)
    add_toc_field(toc)


def add_placeholder(document):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run("[Conținut în lucru]")
    set_font(run, 12, italic=True)


def add_skeleton(document):
    for index, (heading, subheadings) in enumerate(SECTIONS):
        document.add_page_break()
        paragraph = document.add_paragraph(heading, style="Heading 1")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            set_font(run, 14, bold=True)

        add_placeholder(document)

        for subheading in subheadings:
            paragraph = document.add_paragraph(subheading, style="Heading 2")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_font(run, 12, bold=True)
            add_placeholder(document)


def main():
    source = SOURCE if SOURCE.exists() else FALLBACK_SOURCE
    if not source.exists():
        raise FileNotFoundError(source)

    document = Document(source)
    configure_document(document)
    set_update_fields_on_open(document)
    add_toc_page(document)
    add_skeleton(document)

    document.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
