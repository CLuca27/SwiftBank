from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(r"A:\Proiecte\SwiftBank")
OUT_DIR = ROOT / "documentatie"
OUT_PATH = OUT_DIR / "SwiftBank_Coperta_Pagina_Titlu.docx"

UNIVERSITY = "ACADEMIA DE STUDII ECONOMICE DIN BUCUREȘTI"
FACULTY = "Facultatea de Cibernetică, Statistică și Informatică Economică"
SPECIALIZATION = "Specializarea: Informatică economică"
WORK_LABEL = "LUCRARE DE LICENȚĂ"
TITLE = "SISTEM INFORMATIC PENTRU MANAGEMENTUL SERVICIILOR BANCARE DIGITALE"
COORDINATOR = "[Grad didactic, Nume Prenume]"
GRADUATE = "Cirimpei Luca"
CITY_YEAR = "București\n2026"


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = "w:{}".format(edge)
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key, value in edge_data.items():
                element.set(qn("w:{}".format(key)), str(value))


def set_run_font(run, size, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def add_centered_text(document, text, size=12, bold=False, spacing_after=0):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(spacing_after)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return paragraph


def add_spacer(document, points):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(points)
    paragraph.paragraph_format.line_spacing = 1.0
    return paragraph


def add_people_block(document):
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(8)
    table.columns[1].width = Cm(8)

    for cell in table.rows[0].cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_border(
            cell,
            top={"val": "nil"},
            left={"val": "nil"},
            bottom={"val": "nil"},
            right={"val": "nil"},
            insideH={"val": "nil"},
            insideV={"val": "nil"},
        )

    left = table.cell(0, 0)
    right = table.cell(0, 1)

    left_paragraph = left.paragraphs[0]
    left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_paragraph.paragraph_format.line_spacing = 1.0
    run = left_paragraph.add_run("Cadrul didactic coordonator:\n")
    set_run_font(run, 12)
    run = left_paragraph.add_run(COORDINATOR)
    set_run_font(run, 12, bold=True)

    right_paragraph = right.paragraphs[0]
    right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_paragraph.paragraph_format.line_spacing = 1.0
    run = right_paragraph.add_run("Absolvent:\n")
    set_run_font(run, 12)
    run = right_paragraph.add_run(GRADUATE)
    set_run_font(run, 12, bold=True)


def configure_section(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def build_page(document, center_lines):
    add_centered_text(document, UNIVERSITY, size=13, bold=True, spacing_after=8)
    add_centered_text(document, FACULTY, size=12, bold=False, spacing_after=4)
    add_centered_text(document, SPECIALIZATION, size=12, bold=False, spacing_after=0)

    add_spacer(document, 118)

    for index, line in enumerate(center_lines):
        add_centered_text(
            document,
            line,
            size=16 if index == 0 else 14,
            bold=True,
            spacing_after=8,
        )

    add_spacer(document, 150 if len(center_lines) == 1 else 120)
    add_people_block(document)
    add_spacer(document, 52)
    add_centered_text(document, CITY_YEAR, size=12, spacing_after=0)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_section(document.sections[0])

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    build_page(document, [WORK_LABEL])

    section = document.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(section)
    build_page(document, [TITLE])

    document.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
